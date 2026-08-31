# File: backend/app/services/question_import.py
# Purpose: Parse lecturer-uploaded question files (Aiken/GIFT plain text,
#          Excel, Word) into the assignments question schema so the lecturer
#          can review options/answer keys before publishing.
#
# Schema contract: every parser returns
#   ({ "objective": [{"question", "options", "correct_answer_index"}],
#      "theory":    [{"question", "suggested_answer_rubric"}] }, skipped)
# correct_answer_index is the 0-based index of the correct option, or
# -1 when the document does not state a correct answer (lecturer picks in
# the review step). `skipped` is a list of human-readable reasons for the
# lines/rows that could not be imported.

import io
import re

_OPTION_RE = re.compile(r"^([A-Ha-h])[.)]\s*(.+)$")
_ANSWER_RE = re.compile(r"^ANSWER\s*:\s*([A-Ha-h])\s*$", re.IGNORECASE)
_GIFT_BRACES = re.compile(r"\{([^}]*)\}")
_NUMBERED_RE = re.compile(r"^\s*(?:\d{1,2}[.)]|Q\d{1,2}[.)\s:])", re.IGNORECASE)
_LABEL_RE = re.compile(r"^\s*(?:\d+[.)]|Q\d+[.)\s:]*|\(\d+\))[:\-]?\s*", re.IGNORECASE)

_LETTERS = "ABCDEFGH"
_MAX_QUESTIONS = 100


def _strip_label(text: str) -> str:
    s = (text or "").strip()
    s = _LABEL_RE.sub("", s, count=1).strip()
    return s or (text or "").strip()


# ── Plain text (Aiken / GIFT / numbered theory) ────────────────────────

def parse_plain_text(text: str):
    """Parse Aiken, simple GIFT, and numbered theory questions from text."""
    questions = {"objective": [], "theory": []}
    skipped = []
    buf = ""
    options = []
    answer = None

    def flush():
        nonlocal buf, options, answer
        q = _strip_label(buf)
        if options and answer is not None and 0 <= answer < len(options):
            questions["objective"].append(
                {"question": q, "options": options[:], "correct_answer_index": answer}
            )
        elif options and q:
            skipped.append(f"'{q[:48]}…' has options but no ANSWER: line")
        elif q:
            questions["theory"].append({"question": q, "suggested_answer_rubric": ""})
        buf, options, answer = "", [], None

    for raw in (text or "").splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush()
            continue

        am = _ANSWER_RE.match(stripped)
        if am:
            answer = _LETTERS.index(am.group(1).upper())
            continue

        if "{" in stripped and "}" in stripped:
            flush()
            gm = _GIFT_BRACES.search(stripped)
            question_part = _GIFT_BRACES.sub("", stripped).strip()
            content = gm.group(1)
            if not content.strip():
                questions["theory"].append(
                    {"question": _strip_label(question_part), "suggested_answer_rubric": ""}
                )
                continue
            opts = [p for p in content.split("~")]
            clean_opts, correct = [], -1
            for p in opts:
                if p.startswith("="):
                    correct = len(clean_opts)
                    clean_opts.append(p[1:].strip())
                elif p.strip():
                    clean_opts.append(p.lstrip("~").strip())
            if clean_opts:
                questions["objective"].append(
                    {
                        "question": _strip_label(question_part) or "(question)",
                        "options": clean_opts,
                        "correct_answer_index": correct,
                    }
                )
            continue

        om = _OPTION_RE.match(stripped)
        if om:
            if not buf.strip():
                skipped.append(f"'{stripped[:48]}…' — option line with no question before it")
                continue
            options.append(om.group(2).strip())
            continue

        # New question starting while the previous objective is unfinished?
        if options and (re.search(r"[?]$", stripped) or _NUMBERED_RE.match(stripped)):
            flush()

        buf = stripped if not buf else f"{buf} {stripped}"

    flush()
    return _capped(questions), skipped


def _capped(questions):
    return {
        "objective": questions["objective"][:_MAX_QUESTIONS],
        "theory": questions["theory"][:_MAX_QUESTIONS],
    }


# ── Tabular rows (Excel + Word table cells) ────────────────────────────

_HEADER_TOKENS = {
    "question": ("question", "prompt"),
    "answer": ("answer", "correct"),
    "type": ("type", "question type"),
}


def _cval(v):
    if v is None:
        return ""
    return str(v).strip()


def _detect_header(row):
    """Map a header row to {question, options[], answer, type} column indices."""
    if not row:
        return None
    vals = [_cval(v).lower() for v in row]
    if not any(("question" in v or "prompt" in v) for v in vals):
        return None
    mapping = {"question": None, "options": [], "answer": None, "type": None}
    for idx, v in enumerate(vals):
        if mapping["question"] is None and ("question" in v or "prompt" in v or v in ("q", "question text")):
            mapping["question"] = idx
        elif v in ("a", "option a", "choice a") or v.startswith("option") or v.startswith("choice"):
            mapping["options"].append(idx)
        elif v in ("b", "c", "d", "e", "f", "g", "h"):
            mapping["options"].append(idx)
        elif mapping["answer"] is None and ("answer" in v or "correct" in v):
            mapping["answer"] = idx
        elif mapping["type"] is None and v in ("type", "question type", "kind", "format"):
            mapping["type"] = idx
    mapping["options"].sort()
    return mapping


def _resolve_answer(ans_raw, options, row_text):
    """Turn an answer cell (letter / option text / 0-based int) into an index."""
    val = _cval(ans_raw)
    if not val:
        return -1
    one = val[0].upper()
    lowered = val.lower()
    if re.fullmatch(r"[A-H]", one):
        if one in _LETTERS:
            idx = _LETTERS.index(one)
            return idx if idx < len(options) else -1
    if lowered.isdigit():
        n = int(lowered)
        if 1 <= n <= len(options):
            return n - 1  # 1-based in spreadsheets
        if 0 <= n < len(options):
            return n
    for idx, opt in enumerate(options):
        if opt.lower() == lowered:
            return idx
    return -1


def _type_is_objective(type_val, options):
    t = _cval(type_val or "").lower()
    if not t:
        return bool(options)
    return not any(k in t for k in ("theory", "essay", "subjective", "short", "open", "write", "long"))


def parse_rows(rows, skipped):
    """Parse an iterable of row iterables into questions. Used by xlsx + docx tables."""
    questions = {"objective": [], "theory": []}
    rows = list(rows)
    if not rows:
        return questions, skipped

    data = list(rows)
    mapping = _detect_header(data[0])
    start = 1 if mapping else 0
    if not mapping:
        # Positional fallback: Q | A | B | C | D | Answer | Type
        mapping = {
            "question": 0,
            "options": list(range(1, min(5, len(data[0]) if data else 1))),
            "answer": 5,
            "type": 6,
        }

    for row in data[start:]:
        vals = [_cval(v) for v in row]
        q_idx = mapping["question"]
        opt_idxs = mapping["options"]
        q = vals[q_idx] if q_idx is not None and q_idx < len(vals) else ""
        if not q:
            skipped.append("Row skipped: empty question.")
            continue
        options = [vals[i] for i in opt_idxs if i < len(vals) and vals[i]]
        type_val = vals[mapping["type"]] if mapping["type"] is not None and mapping["type"] < len(vals) else ""
        ans_idx = mapping["answer"] if mapping["answer"] is not None else None
        answer = _resolve_answer(vals[ans_idx] if ans_idx is not None and ans_idx < len(vals) else "", options, q) if options else -1

        if _type_is_objective(type_val, options) and options:
            questions["objective"].append(
                {"question": q, "options": options, "correct_answer_index": answer}
            )
        elif not _type_is_objective(type_val, options):
            rubric = ""
            # If the sheet carries a model answer in the answer column, keep it as rubric.
            if ans_idx is not None and ans_idx < len(vals) and _cval(vals[ans_idx]):
                rubric = vals[ans_idx]
            questions["theory"].append({"question": q, "suggested_answer_rubric": rubric})
        else:
            skipped.append(f"'{q[:48]}…' has no answer options.")

    return _capped(questions), skipped


def parse_xlsx(data: bytes):
    """Parse a .xlsx/.xls workbook into questions."""
    import openpyxl

    skipped = []
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    if not wb.sheetnames:
        return {"objective": [], "theory": []}, ["File has no sheets."]
    ws = wb[wb.sheetnames[0]]
    try:
        return parse_rows(ws.iter_rows(values_only=True), skipped)
    except Exception as exc:
        print(f"[question_import] xlsx parse failed: {exc}")
        return {"objective": [], "theory": []}, [f"Could not read the spreadsheet: {exc}"]


def parse_docx(data: bytes):
    """Parse a .docx document: tables (column banks) + paragraphs (Aiken/text)."""
    import docx

    skipped = []
    questions = {"objective": [], "theory": []}

    doc = docx.Document(io.BytesIO(data))
    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            table_rows.append([c.text for c in row.cells])
    if table_rows:
        tq, tskipped = parse_rows(table_rows, skipped)
        questions["objective"].extend(tq["objective"])
        questions["theory"].extend(tq["theory"])
        skipped.extend(tskipped)

    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    if paragraphs:
        pq, pskipped = parse_plain_text("\n".join(paragraphs))
        questions["objective"].extend(pq["objective"])
        questions["theory"].extend(pq["theory"])
        skipped.extend(pskipped)

    return _capped(questions), skipped